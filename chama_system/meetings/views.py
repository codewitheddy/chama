import json
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.messages.views import SuccessMessageMixin
from django.views.generic import ListView, CreateView, UpdateView, DeleteView, DetailView
from django.views import View
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse_lazy
from django.contrib import messages
from django.http import JsonResponse
from .models import Meeting, MeetingAttendance, MeetingPenalty, MeetingPenaltyRule
from .forms import MeetingForm, MeetingMinutesForm, MeetingPenaltyForm, MeetingPenaltyRuleForm
from members.models import Member
from accounts.mixins import TreasurerRequiredMixin, AdminRequiredMixin, MemberAccessMixin


# ── Meetings ─────────────────────────────────────────────────────

class MeetingListView(MemberAccessMixin, ListView):
    model = Meeting
    template_name = 'meetings/meeting_list.html'
    context_object_name = 'meetings'
    paginate_by = 20

    def get_queryset(self):
        qs = super().get_queryset()
        q = self.request.GET.get('q', '').strip()
        status = self.request.GET.get('status', '').strip()
        if q:
            # Search by venue OR date
            from django.db.models import Q
            qs = qs.filter(Q(venue__icontains=q) | Q(date__icontains=q))
        if status:
            qs = qs.filter(status=status)
        return qs

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['q'] = self.request.GET.get('q', '')
        ctx['selected_status'] = self.request.GET.get('status', '')
        # Summary stats
        from django.db.models import Count, Sum
        ctx['total_meetings'] = Meeting.objects.count()
        ctx['held_count'] = Meeting.objects.filter(status='held').count()
        ctx['scheduled_count'] = Meeting.objects.filter(status='scheduled').count()
        ctx['total_penalties_collected'] = (
            MeetingPenalty.objects.aggregate(t=Sum('amount'))['t'] or 0
        )
        return ctx


class MeetingCreateView(TreasurerRequiredMixin, SuccessMessageMixin, CreateView):
    model = Meeting
    form_class = MeetingForm
    template_name = 'meetings/meeting_form.html'
    success_message = "Meeting created."

    def form_valid(self, form):
        response = super().form_valid(form)
        # Pre-populate all members as present so treasurer only marks exceptions
        self.object.auto_populate_attendance()
        return response

    def get_success_url(self):
        return reverse_lazy('meetings:detail', kwargs={'pk': self.object.pk})


class MeetingUpdateView(TreasurerRequiredMixin, SuccessMessageMixin, UpdateView):
    model = Meeting
    form_class = MeetingForm
    template_name = 'meetings/meeting_form.html'
    success_message = "Meeting updated."

    def dispatch(self, request, *args, **kwargs):
        meeting = self.get_object()
        if meeting.status == 'held':
            messages.error(request, "This meeting is held — use 'Edit Minutes' to update minutes only.")
            return redirect('meetings:detail', pk=meeting.pk)
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse_lazy('meetings:detail', kwargs={'pk': self.object.pk})


class MeetingUpdateMinutesView(TreasurerRequiredMixin, SuccessMessageMixin, UpdateView):
    """Allows editing minutes only on a held meeting."""
    model = Meeting
    form_class = MeetingMinutesForm
    template_name = 'meetings/meeting_minutes_form.html'
    success_message = "Minutes updated."

    def get_success_url(self):
        return reverse_lazy('meetings:minutes', kwargs={'pk': self.object.pk})


class MeetingMinutesView(MemberAccessMixin, DetailView):
    """Full-page view/edit for meeting minutes."""
    model = Meeting
    template_name = 'meetings/meeting_minutes_page.html'
    context_object_name = 'meeting'


class MeetingMinutesPdfView(MemberAccessMixin, View):
    """Download meeting minutes as a styled PDF."""

    def get(self, request, pk):
        from io import BytesIO
        from django.http import HttpResponse
        from django.utils import timezone
        from bs4 import BeautifulSoup
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm, mm
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, KeepTogether
        )

        meeting = get_object_or_404(Meeting, pk=pk)

        # ── Colours ───────────────────────────────────────────────
        PRIMARY    = colors.HexColor('#1e3a5f')
        ACCENT     = colors.HexColor('#2563eb')
        GREEN      = colors.HexColor('#16a34a')
        GREEN_BG   = colors.HexColor('#f0fdf4')
        AMBER      = colors.HexColor('#d97706')
        AMBER_BG   = colors.HexColor('#fffbeb')
        RED        = colors.HexColor('#e11d48')
        RED_BG     = colors.HexColor('#fff1f2')
        BLUE_BG    = colors.HexColor('#eff6ff')
        ALT_ROW    = colors.HexColor('#f8fafc')
        BORDER     = colors.HexColor('#e2e8f0')
        HEADER_BG  = colors.HexColor('#f1f5f9')
        TEXT_DARK  = colors.HexColor('#1e293b')
        TEXT_MUTED = colors.HexColor('#64748b')
        WHITE      = colors.white

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="minutes_{meeting.date}.pdf"'
        buffer = BytesIO()

        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            topMargin=34*mm, bottomMargin=18*mm,
            leftMargin=1.8*cm, rightMargin=1.8*cm,
        )
        usable = A4[0] - 3.6*cm
        styles = getSampleStyleSheet()

        # ── Page header/footer ────────────────────────────────────
        def on_page(canvas, d):
            canvas.saveState()
            w, h = A4
            # Header bar
            canvas.setFillColor(PRIMARY)
            canvas.rect(0, h - 28*mm, w, 28*mm, fill=1, stroke=0)
            canvas.setFillColor(WHITE)
            canvas.setFont('Helvetica-Bold', 13)
            canvas.drawString(1.8*cm, h - 13*mm, 'DC Welfare Group')
            canvas.setFont('Helvetica', 8)
            canvas.drawString(1.8*cm, h - 21*mm, 'Meeting Minutes')
            canvas.setFont('Helvetica', 9)
            canvas.drawRightString(w - 1.8*cm, h - 13*mm, str(meeting.date))
            canvas.setFont('Helvetica', 8)
            canvas.drawRightString(w - 1.8*cm, h - 21*mm,
                                   f'Generated: {timezone.localdate().strftime("%d %B %Y")}')
            # Footer bar
            canvas.setFillColor(BORDER)
            canvas.rect(0, 0, w, 10*mm, fill=1, stroke=0)
            canvas.setFillColor(TEXT_MUTED)
            canvas.setFont('Helvetica', 8)
            canvas.drawString(1.8*cm, 3.5*mm, 'DC Welfare Group — Confidential')
            canvas.drawRightString(w - 1.8*cm, 3.5*mm, f'Page {d.page}')
            canvas.restoreState()

        # ── Paragraph styles ──────────────────────────────────────
        def ps(name, **kw):
            return ParagraphStyle(name, parent=styles['Normal'], **kw)

        title_s    = ps('title',   fontSize=20, fontName='Helvetica-Bold',
                         textColor=PRIMARY, alignment=TA_CENTER, spaceAfter=2*mm)
        meta_s     = ps('meta',    fontSize=9,  fontName='Helvetica',
                         textColor=TEXT_MUTED, alignment=TA_CENTER, spaceAfter=1*mm)
        section_s  = ps('sec',     fontSize=11, fontName='Helvetica-Bold',
                         textColor=PRIMARY, spaceBefore=5*mm, spaceAfter=2*mm)
        cell_s     = ps('cell',    fontSize=8,  fontName='Helvetica',
                         textColor=TEXT_DARK, leading=12)
        bold_s     = ps('bold',    fontSize=8,  fontName='Helvetica-Bold',
                         textColor=TEXT_DARK, leading=12)
        muted_s    = ps('muted',   fontSize=8,  fontName='Helvetica',
                         textColor=TEXT_MUTED, leading=12)
        hdr_s      = ps('hdr',     fontSize=8,  fontName='Helvetica-Bold',
                         textColor=WHITE, leading=12)
        # Minutes content styles
        body_s     = ps('body',    fontSize=10, fontName='Helvetica',
                         textColor=TEXT_DARK, leading=16, spaceAfter=4*mm)
        h1_s       = ps('h1',      fontSize=14, fontName='Helvetica-Bold',
                         textColor=PRIMARY, spaceBefore=6*mm, spaceAfter=2*mm)
        h2_s       = ps('h2',      fontSize=12, fontName='Helvetica-Bold',
                         textColor=PRIMARY, spaceBefore=4*mm, spaceAfter=2*mm)
        h3_s       = ps('h3',      fontSize=10, fontName='Helvetica-Bold',
                         textColor=TEXT_DARK, spaceBefore=3*mm, spaceAfter=1*mm)
        li_s       = ps('li',      fontSize=10, fontName='Helvetica',
                         textColor=TEXT_DARK, leading=16, leftIndent=12,
                         spaceAfter=2*mm)
        quote_s    = ps('quote',   fontSize=10, fontName='Helvetica-Oblique',
                         textColor=TEXT_MUTED, leading=16, leftIndent=16,
                         borderPadding=(4, 8, 4, 8), spaceAfter=4*mm)

        def make_table(data, col_ratios, extra_styles=None):
            widths = [usable * r for r in col_ratios]
            t = Table(data, colWidths=widths, repeatRows=1)
            cmds = [
                ('BACKGROUND',     (0, 0), (-1, 0),  PRIMARY),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [WHITE, ALT_ROW]),
                ('GRID',           (0, 0), (-1, -1), 0.4, BORDER),
                ('LINEBELOW',      (0, 0), (-1, 0),  1.2, ACCENT),
                ('TOPPADDING',     (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING',  (0, 0), (-1, -1), 5),
                ('LEFTPADDING',    (0, 0), (-1, -1), 6),
                ('RIGHTPADDING',   (0, 0), (-1, -1), 6),
                ('VALIGN',         (0, 0), (-1, -1), 'MIDDLE'),
            ]
            if extra_styles:
                cmds.extend(extra_styles)
            t.setStyle(TableStyle(cmds))
            return t

        elements = []

        # ── Title block ───────────────────────────────────────────
        elements.append(Paragraph('Meeting Minutes', title_s))
        elements.append(Paragraph(
            f'Date: <b>{meeting.date}</b> &nbsp;|&nbsp; '
            f'Venue: <b>{meeting.venue or "—"}</b> &nbsp;|&nbsp; '
            f'Status: <b>{meeting.get_status_display()}</b>',
            meta_s
        ))
        elements.append(HRFlowable(
            width='100%', thickness=2, color=ACCENT,
            spaceBefore=3*mm, spaceAfter=4*mm
        ))

        # ── Agenda ────────────────────────────────────────────────
        if meeting.agenda:
            elements.append(Paragraph('Agenda', section_s))
            for line in meeting.agenda.strip().splitlines():
                line = line.strip()
                if line:
                    elements.append(Paragraph(f'• {line}', li_s))
            elements.append(Spacer(1, 2*mm))

        # ── Attendance ────────────────────────────────────────────
        attendance = meeting.attendance.select_related('member').order_by('member__name')
        if attendance.exists():
            elements.append(Paragraph('Attendance', section_s))

            # Summary counts
            from collections import Counter
            counts = Counter(a.status for a in attendance)
            summary_data = [[
                Paragraph('<b>Present</b>', bold_s),
                Paragraph('<b>Late</b>', bold_s),
                Paragraph('<b>Apology</b>', bold_s),
                Paragraph('<b>Absent</b>', bold_s),
                Paragraph('<b>Total</b>', bold_s),
            ], [
                Paragraph(str(counts.get('present', 0)), cell_s),
                Paragraph(str(counts.get('late', 0)), cell_s),
                Paragraph(str(counts.get('absent_apology', 0)), cell_s),
                Paragraph(str(counts.get('absent_no_apology', 0)), cell_s),
                Paragraph(str(attendance.count()), cell_s),
            ]]
            summary_tbl = Table(summary_data, colWidths=[usable/5]*5)
            summary_tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0, 0), (-1, 0),  HEADER_BG),
                ('GRID',          (0, 0), (-1, -1), 0.4, BORDER),
                ('TOPPADDING',    (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING',   (0, 0), (-1, -1), 8),
                ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(summary_tbl)
            elements.append(Spacer(1, 3*mm))

            # Attendance register
            att_data = [[
                Paragraph('<b>#</b>', hdr_s),
                Paragraph('<b>MEMBER</b>', hdr_s),
                Paragraph('<b>STATUS</b>', hdr_s),
            ]]
            status_colors = {
                'present':           (GREEN,   'Present'),
                'late':              (AMBER,   'Late'),
                'absent_apology':    (ACCENT,  'Apology'),
                'absent_no_apology': (RED,     'Absent'),
            }
            row_extra = []
            for i, a in enumerate(attendance, 1):
                color, label = status_colors.get(a.status, (TEXT_MUTED, a.get_status_display()))
                status_p = Paragraph(label, ps(f'st{a.pk}', fontSize=8,
                                               fontName='Helvetica-Bold',
                                               textColor=color, leading=12))
                att_data.append([
                    Paragraph(str(i), muted_s),
                    Paragraph(a.member.name, cell_s),
                    status_p,
                ])
            elements.append(make_table(att_data, [0.08, 0.62, 0.30]))
            elements.append(Spacer(1, 2*mm))

        # ── Penalties ─────────────────────────────────────────────
        penalties = meeting.penalties.select_related('member', 'rule').all()
        if penalties.exists():
            elements.append(Paragraph('Penalties Issued', section_s))
            pen_data = [[
                Paragraph('<b>MEMBER</b>', hdr_s),
                Paragraph('<b>REASON</b>', hdr_s),
                Paragraph('<b>AMOUNT (KES)</b>', hdr_s),
            ]]
            total_pen = 0
            for p in penalties:
                pen_data.append([
                    Paragraph(p.member.name, cell_s),
                    Paragraph(p.reason, cell_s),
                    Paragraph(f'{p.amount:,.2f}', ps(f'pa{p.pk}', fontSize=8,
                                                      fontName='Helvetica',
                                                      textColor=RED, leading=12,
                                                      alignment=TA_RIGHT)),
                ])
                total_pen += p.amount
            pen_data.append([
                Paragraph(''), Paragraph('<b>Total</b>', bold_s),
                Paragraph(f'<b>{total_pen:,.2f}</b>', ps('ptot', fontSize=8,
                                                          fontName='Helvetica-Bold',
                                                          textColor=RED, leading=12,
                                                          alignment=TA_RIGHT)),
            ])
            pen_tbl = make_table(pen_data, [0.35, 0.45, 0.20],
                                 extra_styles=[('BACKGROUND', (0, -1), (-1, -1), HEADER_BG)])
            elements.append(pen_tbl)
            elements.append(Spacer(1, 2*mm))

        # ── Minutes content ───────────────────────────────────────
        elements.append(HRFlowable(
            width='100%', thickness=1, color=BORDER,
            spaceBefore=4*mm, spaceAfter=4*mm
        ))
        elements.append(Paragraph('Minutes', section_s))

        if meeting.minutes:
            soup = BeautifulSoup(meeting.minutes, 'html.parser')

            def safe_para(text, style):
                """Return a Paragraph, falling back to plain text if markup is broken."""
                if not text or not text.strip():
                    return None
                # Strip all HTML tags — ReportLab's parser is strict
                from html.parser import HTMLParser
                class Stripper(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.parts = []
                    def handle_data(self, d):
                        self.parts.append(d)
                    def get(self):
                        return ''.join(self.parts)
                s = Stripper()
                s.feed(text)
                clean = s.get().strip()
                if not clean:
                    return None
                try:
                    return Paragraph(clean, style)
                except Exception:
                    return Paragraph(clean.replace('<', '&lt;').replace('>', '&gt;'), style)

            for el in soup.find_all(
                ['h1', 'h2', 'h3', 'p', 'li', 'blockquote', 'ul', 'ol'],
                recursive=True
            ):
                tag = el.name
                # Skip nested li inside ul/ol — we handle them via ul/ol
                if tag in ('ul', 'ol'):
                    continue
                text = el.get_text()
                if not text.strip():
                    continue
                if tag == 'h1':
                    p = safe_para(text, h1_s)
                elif tag == 'h2':
                    p = safe_para(text, h2_s)
                elif tag == 'h3':
                    p = safe_para(text, h3_s)
                elif tag == 'li':
                    p = safe_para(f'• {text}', li_s)
                elif tag == 'blockquote':
                    p = safe_para(text, quote_s)
                else:
                    p = safe_para(text, body_s)
                if p:
                    elements.append(p)
        else:
            elements.append(Paragraph('No minutes recorded.', muted_s))

        # ── Signature block ───────────────────────────────────────
        elements.append(Spacer(1, 12*mm))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=BORDER, spaceAfter=4*mm))
        sig_data = [[
            Paragraph('Chairperson\n\n\n_______________________', ps('sig', fontSize=9,
                       fontName='Helvetica', textColor=TEXT_MUTED, leading=14)),
            Paragraph('Secretary\n\n\n_______________________', ps('sig2', fontSize=9,
                       fontName='Helvetica', textColor=TEXT_MUTED, leading=14)),
            Paragraph('Treasurer\n\n\n_______________________', ps('sig3', fontSize=9,
                       fontName='Helvetica', textColor=TEXT_MUTED, leading=14)),
        ]]
        sig_tbl = Table(sig_data, colWidths=[usable/3]*3)
        sig_tbl.setStyle(TableStyle([
            ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
            ('TOPPADDING',    (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        elements.append(sig_tbl)

        doc.build(elements, onFirstPage=on_page, onLaterPages=on_page)
        response.write(buffer.getvalue())
        buffer.close()
        return response


class MeetingMinutesDocxView(MemberAccessMixin, View):
    """Download minutes as a .docx file."""
    def get(self, request, pk):
        from docx import Document
        from docx.shared import Pt, RGBColor
        from bs4 import BeautifulSoup
        from django.http import HttpResponse
        import io

        meeting = get_object_or_404(Meeting, pk=pk)
        doc = Document()

        # Title
        title = doc.add_heading(f'Meeting Minutes — {meeting.date}', 0)
        title.alignment = 1  # center

        # Meta info
        doc.add_paragraph(f'Venue: {meeting.venue or "—"}')
        doc.add_paragraph(f'Status: {meeting.get_status_display()}')
        doc.add_paragraph(f'Attendance: {meeting.attendance.count()} members')
        doc.add_paragraph('')

        # Parse HTML minutes
        if meeting.minutes:
            soup = BeautifulSoup(meeting.minutes, 'html.parser')
            for el in soup.find_all(['p', 'h1', 'h2', 'h3', 'li', 'blockquote']):
                tag = el.name
                text = el.get_text()
                if not text.strip():
                    continue
                if tag in ('h1', 'h2'):
                    doc.add_heading(text, level=1)
                elif tag == 'h3':
                    doc.add_heading(text, level=2)
                elif tag == 'li':
                    doc.add_paragraph(text, style='List Bullet')
                elif tag == 'blockquote':
                    p = doc.add_paragraph(text)
                    p.style = 'Intense Quote'
                else:
                    doc.add_paragraph(text)
        else:
            doc.add_paragraph('No minutes recorded.')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f'minutes_{meeting.date}.docx'
        response = HttpResponse(
            buf,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response


class MeetingDeleteView(AdminRequiredMixin, DeleteView):
    model = Meeting
    template_name = 'meetings/meeting_confirm_delete.html'
    success_url = reverse_lazy('meetings:list')


class MeetingDetailView(MemberAccessMixin, DetailView):
    model = Meeting
    template_name = 'meetings/meeting_detail.html'
    context_object_name = 'meeting'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        meeting = self.object
        attendance = meeting.attendance.select_related('member').all()
        attended_ids = attendance.values_list('member_id', flat=True)
        ctx['attendance'] = attendance
        ctx['penalties'] = meeting.penalties.select_related('member', 'rule').all()
        ctx['penalty_form'] = MeetingPenaltyForm(meeting=meeting)
        ctx['rules'] = list(MeetingPenaltyRule.objects.values('id', 'name', 'default_amount'))
        ctx['all_members'] = Member.objects.all()
        ctx['unrecorded_members'] = Member.objects.exclude(pk__in=attended_ids)
        return ctx


# ── Attendance ────────────────────────────────────────────────────

class SaveAttendanceView(TreasurerRequiredMixin, View):
    """Bulk-save attendance for a meeting."""
    def post(self, request, pk):
        from django.db import transaction
        meeting = get_object_or_404(Meeting, pk=pk)
        if meeting.status == 'held':
            return JsonResponse({'ok': False, 'error': 'Meeting is held — attendance is locked.'}, status=403)
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'ok': False, 'error': 'Invalid JSON.'}, status=400)

        if not isinstance(data, list):
            return JsonResponse({'ok': False, 'error': 'Expected a list payload.'}, status=400)

        valid_statuses = {choice[0] for choice in MeetingAttendance.STATUS_CHOICES}
        member_ids = [row.get('member_id') for row in data if isinstance(row, dict)]
        existing_member_ids = set(
            Member.objects.filter(pk__in=member_ids).values_list('pk', flat=True)
        )

        invalid_rows = []
        with transaction.atomic():
            for row in data:
                if not isinstance(row, dict):
                    invalid_rows.append('Invalid row format.')
                    continue
                member_id = row.get('member_id')
                status = row.get('status')
                if member_id not in existing_member_ids:
                    invalid_rows.append(f'Unknown member_id: {member_id}')
                    continue
                if status not in valid_statuses:
                    invalid_rows.append(f'Invalid attendance status: {status}')
                    continue
                MeetingAttendance.objects.update_or_create(
                    meeting=meeting,
                    member_id=member_id,
                    defaults={'status': status},
                )

        if invalid_rows:
            return JsonResponse({'ok': False, 'error': invalid_rows[0]}, status=400)

        return JsonResponse({'ok': True, 'saved': len(data) - len(invalid_rows)})


# ── Penalties ─────────────────────────────────────────────────────

class AddMeetingPenaltyView(TreasurerRequiredMixin, View):
    def post(self, request, pk):
        meeting = get_object_or_404(Meeting, pk=pk)
        if meeting.status == 'held':
            messages.error(request, "Penalties cannot be added once a meeting is marked as held.")
            return redirect('meetings:detail', pk=pk)
        form = MeetingPenaltyForm(request.POST, meeting=meeting)
        if form.is_valid():
            mp = form.save(commit=False)
            mp.meeting = meeting
            mp.save()
            messages.success(request, f"Penalty of KES {mp.amount} recorded for {mp.member.name}.")
        else:
            for err in form.errors.values():
                messages.error(request, err.as_text())
        return redirect('meetings:detail', pk=pk)


class ApplyAbsentPenaltiesView(TreasurerRequiredMixin, View):
    """
    Auto-apply the selected penalty rule to all members marked
    'absent_no_apology' who don't already have a penalty for this meeting.
    """
    def post(self, request, pk):
        from django.db import transaction
        meeting = get_object_or_404(Meeting, pk=pk)
        if meeting.status == 'held':
            messages.error(request, "Meeting is held — penalties are locked.")
            return redirect('meetings:detail', pk=pk)

        rule_id = request.POST.get('rule_id')
        if not rule_id:
            messages.error(request, "Please select a penalty rule.")
            return redirect('meetings:detail', pk=pk)

        try:
            rule = MeetingPenaltyRule.objects.get(pk=rule_id)
        except MeetingPenaltyRule.DoesNotExist:
            messages.error(request, "Penalty rule not found.")
            return redirect('meetings:detail', pk=pk)

        # Members absent without apology
        absent_ids = meeting.attendance.filter(
            status='absent_no_apology'
        ).values_list('member_id', flat=True)

        # Exclude those who already have a penalty for this meeting
        already_penalised = meeting.penalties.values_list('member_id', flat=True)
        to_penalise = Member.objects.filter(
            pk__in=absent_ids
        ).exclude(pk__in=already_penalised)

        count = 0
        with transaction.atomic():
            for member in to_penalise:
                MeetingPenalty.objects.create(
                    meeting=meeting,
                    member=member,
                    rule=rule,
                    amount=rule.default_amount,
                    reason=f"Absent without apology — {rule.name}",
                )
                count += 1

        if count:
            messages.success(request, f"Applied {rule.name} penalty (KES {rule.default_amount}) to {count} absent member(s).")
        else:
            messages.info(request, "No new absent members to penalise.")
        return redirect('meetings:detail', pk=pk)


class DeleteMeetingPenaltyView(AdminRequiredMixin, View):
    def post(self, request, pk):
        mp = get_object_or_404(MeetingPenalty, pk=pk)
        if mp.meeting.status == 'held':
            messages.error(request, "Penalties cannot be deleted once a meeting is marked as held.")
            return redirect('meetings:detail', pk=mp.meeting_id)
        meeting_pk = mp.meeting_id
        mp.delete()
        messages.success(request, "Penalty removed.")
        return redirect('meetings:detail', pk=meeting_pk)


# ── Penalty Rules ─────────────────────────────────────────────────

class PenaltyRuleListView(AdminRequiredMixin, ListView):
    model = MeetingPenaltyRule
    template_name = 'meetings/penalty_rules.html'
    context_object_name = 'rules'


class PenaltyRuleCreateView(AdminRequiredMixin, SuccessMessageMixin, CreateView):
    model = MeetingPenaltyRule
    form_class = MeetingPenaltyRuleForm
    template_name = 'meetings/penalty_rule_form.html'
    success_url = reverse_lazy('meetings:rules')
    success_message = "Penalty rule created."


class PenaltyRuleUpdateView(AdminRequiredMixin, SuccessMessageMixin, UpdateView):
    model = MeetingPenaltyRule
    form_class = MeetingPenaltyRuleForm
    template_name = 'meetings/penalty_rule_form.html'
    success_url = reverse_lazy('meetings:rules')
    success_message = "Penalty rule updated."


class PenaltyRuleDeleteView(AdminRequiredMixin, DeleteView):
    model = MeetingPenaltyRule
    template_name = 'meetings/penalty_rule_confirm_delete.html'
    success_url = reverse_lazy('meetings:rules')
